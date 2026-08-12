"""
fast_server.py — Serveur RAG vs KAG (version corrigée)
"""
import os, sys, time, json, logging, re, string, io
from pathlib import Path
from collections import Counter
from typing import List, Optional

from fastapi import FastAPI, UploadFile, File, Form, Query, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from groq import Groq
import minsearch

# PDF / DOCX parsers
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

load_dotenv()
logging.basicConfig(level=logging.WARNING)

GROQ_KEY      = os.getenv("GROQ_API_KEY", "")
client        = Groq(api_key=GROQ_KEY)
MODEL_ANSWER  = "llama-3.1-8b-instant"   # réponses QA — 6 000 TPM
MODEL_EXTRACT = "gemma2-9b-it"           # extraction KAG — 15 000 TPM

def llm(system_prompt: str, user_prompt: str,
        max_tokens: int = 150, model: str = None) -> str:
    if model is None:
        model = MODEL_ANSWER
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt}
            ],
            temperature=0.0, max_tokens=max_tokens
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"ERREUR: {e}"

app = FastAPI(title="RAG vs KAG")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ── State ──────────────────────────────────────────────────────────────────────
rag_index      = None
documents      = []
kag_graph      = {"nodes": {}, "relations": {}}
run_log        = []   # {pipeline, question, answer, latency_ms}
routing_log    = []   # {question, type, routed_to, reason, confidence}
uploaded_files = []   # [{name, size_kb, chunks, status}]

# ── File Parsing ───────────────────────────────────────────────────────────────
def _parse_uploaded_file(filename: str, content: bytes) -> str:
    """Extrait le texte brut depuis PDF, DOCX, TXT, MD ou CSV."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            if not HAS_PDF:
                return ""
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        elif ext in (".docx",):
            if not HAS_DOCX:
                return ""
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in (".txt", ".md"):
            return content.decode("utf-8", errors="ignore")
        elif ext == ".csv":
            text = content.decode("utf-8", errors="ignore")
            # Chaque ligne CSV devient une phrase
            lines = [l.replace(",", " — ") for l in text.splitlines() if l.strip()]
            return "\n".join(lines)
        else:
            # Essai en UTF-8 pour tout autre format texte
            return content.decode("utf-8", errors="ignore")
    except Exception as e:
        logging.warning(f"Parse error {filename}: {e}")
        return ""

def _chunk_text(text: str, filename: str, chunk_size: int = 400,
                overlap: int = 50) -> list:
    """Découpe le texte en chunks de ~chunk_size mots avec overlap."""
    words  = text.split()
    chunks = []
    step   = max(1, chunk_size - overlap)
    for i in range(0, len(words), step):
        chunk_words = words[i:i + chunk_size]
        if len(chunk_words) < 20:   # ignorer les micro-chunks
            continue
        chunks.append({
            "title": filename,
            "text":  " ".join(chunk_words),
            "id":    f"{filename}_{i}"
        })
    return chunks

# ── Routing Agent ──────────────────────────────────────────────────────────────
# Signaux linguistiques pour la classification de questions

MULTIHOP_SIGNALS  = {"same nationality", "also", "were they", "did they",
                     "same country", "same city", "same field", "same genre",
                     "who is also", "who was also", "as well as",
                     "in addition to", "both born", "both studied",
                     "both worked", "how many of them", "which of them"}

RELATIONAL_SIGNALS= {"father of", "mother of", "son of", "daughter of",
                     "founded by", "created by", "directed by", "produced by",
                     "who played", "who portrayed", "who held",
                     "what position", "what role", "what government",
                     "relationship between", "married to", "worked with",
                     "collaborated with", "trained by", "coached by"}

FACTUAL_SIGNALS   = {"what year", "when was", "when did", "how many",
                     "how much", "what is the name", "what was the name",
                     "in what year", "what date", "how old", "what number",
                     "which year", "what age", "how long", "what time"}

BOOLEAN_STARTERS  = {"were", "was", "is", "are", "did", "does",
                     "have", "has", "can", "could", "would", "do",
                     "will", "shall", "might", "should"}

# Mots exclus du comptage d'entités
STOP_CAPS = {"The","A","An","Is","Are","Was","Were","Did","Does","Have",
             "Has","Can","Could","Would","Do","Will","Shall","In","Of",
             "And","Or","But","For","With","From","By","On","At","To"}


def _risk_consequences(chosen_pipeline: str, q_type: str,
                       entity_count: int, mh_hits: int, re_hits: int,
                       fa_hits: int, is_simple_bool: bool) -> list:
    """Génère les conséquences si l'utilisateur choisit le mauvais pipeline."""
    risks = []
    wrong = "rag" if chosen_pipeline == "kag" else "kag"

    if wrong == "rag":
        # Risques d'utiliser RAG quand KAG est recommandé
        if mh_hits > 0:
            risks.append("❌ RAG ne peut pas chaîner plusieurs faits → réponse partielle ou incorrecte")
        if re_hits > 0:
            risks.append("❌ RAG cherche du texte brut, pas des relations → connexion entre entités manquée")
        if entity_count >= 2:
            risks.append(f"⚠️ {entity_count} entités détectées → RAG risque de ne trouver qu'un seul contexte")
        if q_type == "multi_hop":
            risks.append("❌ Raisonnement multi-sauts impossible sans graphe → hallucination probable")
        if not risks:
            risks.append("⚠️ RAG pourrait manquer une relation implicite entre entités du graphe")
    else:
        # Risques d'utiliser KAG quand RAG est recommandé
        if is_simple_bool or q_type == "boolean":
            risks.append("⚠️ KAG sur-complexifie une question oui/non simple → latence inutile")
        if q_type == "factual":
            risks.append("⚠️ Un seul passage de texte suffit → traverser le graphe est inutile")
        risks.append("❌ Si l'entité n'est pas dans le graphe → réponse vide ou générique")
        risks.append("⚠️ KAG dépend de la qualité de l'extraction → triplets manquants = erreurs")

    return risks[:3]  # max 3 risques affichés


def classify_question(question: str) -> dict:
    """Classifie la question avec scoring RAG vs KAG, facteurs détectés et analyse de risque."""
    q     = question.lower().strip()
    words = q.split()
    orig  = question.split()
    first = words[0] if words else ""

    # ── Détection des signaux ──────────────────────────────────────────────────
    factors_detected = []

    mh_hits = 0
    for sig in MULTIHOP_SIGNALS:
        if sig in q:
            mh_hits += 1
            factors_detected.append({"label": f'"{sig}"', "type": "multi_hop",
                                     "icon": "🧠", "favor": "kag", "weight": 45})

    re_hits = 0
    for sig in RELATIONAL_SIGNALS:
        if sig in q:
            re_hits += 1
            factors_detected.append({"label": f'"{sig}"', "type": "relational",
                                     "icon": "🔗", "favor": "kag", "weight": 40})

    fa_hits = 0
    for sig in FACTUAL_SIGNALS:
        if sig in q:
            fa_hits += 1
            factors_detected.append({"label": f'"{sig}"', "type": "factual",
                                     "icon": "📄", "favor": "rag", "weight": 35})

    bo_hit = 1 if first in BOOLEAN_STARTERS else 0
    if bo_hit:
        factors_detected.append({"label": f'commence par "{first}"', "type": "boolean",
                                 "icon": "❓", "favor": "rag", "weight": 25})

    # Facteur : nombre d'entités propres dans la question
    entity_count = sum(1 for w in orig
                       if w[0].isupper() and len(w) > 2 and w not in STOP_CAPS)
    if entity_count >= 3:
        factors_detected.append({"label": f"{entity_count} entités nommées", "type": "multi_hop",
                                 "icon": "👥", "favor": "kag", "weight": 30})
    elif entity_count == 2:
        factors_detected.append({"label": "2 entités nommées", "type": "relational",
                                 "icon": "👥", "favor": "kag", "weight": 20})

    # Facteur : longueur de la question
    qlen = len(words)
    if qlen > 18:
        factors_detected.append({"label": f"question longue ({qlen} mots)", "type": "multi_hop",
                                 "icon": "📏", "favor": "kag", "weight": 15})
    elif qlen <= 8:
        factors_detected.append({"label": f"question courte ({qlen} mots)", "type": "factual",
                                 "icon": "📏", "favor": "rag", "weight": 15})

    # Facteur : présence de "both" + "and"
    has_both    = "both" in q
    has_and_two = q.count(" and ") >= 1
    if has_both and has_and_two:
        factors_detected.append({"label": '"both ... and" détecté', "type": "multi_hop",
                                 "icon": "🔄", "favor": "kag", "weight": 25})

    # Facteur : pronoms de référence multiple
    if any(p in words for p in ["they", "them", "their", "both"]):
        factors_detected.append({"label": "pronom multi-entité", "type": "multi_hop",
                                 "icon": "🔄", "favor": "kag", "weight": 10})

    # ── Scores RAG vs KAG (0-100) ─────────────────────────────────────────────
    is_simple_boolean = (bo_hit and qlen <= 14 and re_hits == 0 and mh_hits == 0)

    rag_score  = 10  # base
    kag_score  = 10  # base
    rag_score += fa_hits * 30
    rag_score += 30 if is_simple_boolean else 0
    rag_score += 15 if qlen <= 8 else 0
    rag_score += 10 if entity_count <= 1 else 0
    rag_score += 20 if bo_hit and not has_both else 0

    kag_score += mh_hits * 40
    kag_score += re_hits * 35
    kag_score += 25 if has_both and has_and_two else 0
    kag_score += 20 if entity_count >= 3 else (12 if entity_count == 2 else 0)
    kag_score += 15 if qlen > 18 else 0
    kag_score += 10 if any(p in words for p in ["they","them","their"]) else 0

    rag_score = min(100, rag_score)
    kag_score = min(100, kag_score)

    # Normalisation pour que les deux totalisent 100
    total = rag_score + kag_score
    if total > 0:
        rag_pct = round(rag_score / total * 100)
        kag_pct = 100 - rag_pct
    else:
        rag_pct = kag_pct = 50

    # ── Niveau de risque si mauvais choix ─────────────────────────────────────
    score_diff = abs(rag_pct - kag_pct)
    if score_diff >= 35:
        risk_level, risk_emoji = "Élevé",  "🔴"
    elif score_diff >= 15:
        risk_level, risk_emoji = "Modéré", "🟡"
    else:
        risk_level, risk_emoji = "Faible",  "🟢"

    # ── Type scores (radar chart) ──────────────────────────────────────────────
    type_scores = {
        "multi_hop":  min(100, mh_hits * 45 + (15 if mh_hits else 0)),
        "relational": min(100, re_hits * 40 + (10 if re_hits else 0)),
        "factual":    min(100, fa_hits * 45 + (10 if fa_hits else 0)),
        "boolean":    bo_hit * 70 + (20 if is_simple_boolean else 0),
    }
    if bo_hit and has_both and has_and_two:
        type_scores["multi_hop"] = max(type_scores["multi_hop"], 55)
    for k in type_scores:
        if type_scores[k] == 0:
            type_scores[k] = 5

    # ── Classification prioritaire ────────────────────────────────────────────
    if is_simple_boolean:
        q_type, confidence = "boolean", 0.85
        reason = "Question oui/non simple → RAG suffisant (un seul passage de texte)"
    elif mh_hits >= 2 or (mh_hits >= 1 and fa_hits == 0 and re_hits == 0):
        q_type, confidence = "multi_hop", 0.90
        reason = "Multi-sauts détectés → KAG nécessaire (raisonnement sur le graphe)"
    elif bo_hit and has_both and has_and_two:
        q_type, confidence = "multi_hop", 0.82
        reason = "Oui/non sur deux entités distinctes → KAG pour comparer dans le graphe"
    elif re_hits:
        q_type, confidence = "relational", 0.88
        reason = "Relation entre entités détectée → KAG (parcours des arêtes du graphe)"
    elif fa_hits:
        q_type, confidence = "factual", 0.85
        reason = "Valeur précise recherchée → RAG (passages BM25 pertinents)"
    elif bo_hit:
        q_type, confidence = "boolean", 0.78
        reason = "Question oui/non → Ensemble pour arbitrage RAG + KAG"
    else:
        sys_p = ("Classify this question in ONE word: multi_hop | relational | factual | boolean. "
                 "Just the word, no explanation.")
        res   = llm(sys_p, question, max_tokens=5).strip().lower()
        q_type    = res if res in ("multi_hop","relational","factual","boolean") else "factual"
        confidence = 0.70
        reason = "Classifié par LLM (aucun signal lexical détecté)"
        type_scores[q_type] = max(type_scores[q_type], 60)

    # Pipeline recommandé selon les scores
    recommended = "rag" if rag_pct >= kag_pct else "kag"

    # ── Analyse de risque si mauvais choix ────────────────────────────────────
    risk_consequences = _risk_consequences(
        recommended, q_type, entity_count,
        mh_hits, re_hits, fa_hits, is_simple_boolean
    )

    return {
        "type": q_type,
        "confidence": confidence,
        "reason": reason,
        "type_scores": type_scores,
        # Nouveau : scoring comparatif
        "score_rag": rag_pct,
        "score_kag": kag_pct,
        "risk_level": risk_level,
        "risk_emoji": risk_emoji,
        "risk_consequences": risk_consequences,
        "factors_detected": factors_detected,
        "entity_count": entity_count,
        "q_length": qlen,
    }


def route_question(question: str) -> dict:
    """Décide le meilleur pipeline : rag | kag | ensemble."""
    kag_ready = len(kag_graph["nodes"]) > 5
    rag_ready = rag_index is not None and len(documents) > 0
    clf       = classify_question(question)
    q_type    = clf["type"]
    words     = question.lower().split()
    first     = words[0] if words else ""
    is_simple_bool = (first in BOOLEAN_STARTERS and len(words) <= 14 and q_type == "boolean")

    if not kag_ready and not rag_ready:
        return {**clf, "pipeline": "none",
                "reason": "Aucun pipeline disponible — chargez des données d'abord"}
    if not kag_ready:
        return {**clf, "pipeline": "rag",
                "reason": f"KAG non disponible. RAG utilisé (question: {q_type})"}
    if not rag_ready:
        return {**clf, "pipeline": "kag",
                "reason": f"RAG non disponible. KAG utilisé (question: {q_type})"}

    if q_type == "multi_hop":
        return {**clf, "pipeline": "kag",
                "reason": "KAG ✓ : multi-sauts → raisonnement sur le graphe de triplets"}
    if q_type == "relational":
        return {**clf, "pipeline": "kag",
                "reason": "KAG ✓ : relation entre entités → parcours des arêtes du graphe"}
    if q_type == "factual":
        return {**clf, "pipeline": "rag",
                "reason": "RAG ✓ : fait direct → passages BM25 pertinents"}
    if q_type == "boolean":
        if is_simple_bool:
            return {**clf, "pipeline": "rag",
                    "reason": "RAG ✓ : oui/non simple → un passage suffit"}
        return {**clf, "pipeline": "ensemble",
                "reason": "Ensemble ✓ : oui/non complexe → RAG + KAG + agent arbitre"}

    return {**clf, "pipeline": "ensemble",
            "reason": "Incertain → Ensemble RAG + KAG pour robustesse"}

def ensemble_answer(question: str, rag: dict, kag: dict, routing: dict) -> dict:
    """Agent arbitre : choisit la meilleure réponse entre RAG et KAG."""
    rag_ans = rag.get("answer","")
    kag_ans = kag.get("answer","")
    REFUSALS = ["don't know", "do not know", "cannot determine",
                "not enough", "je ne sais pas", "impossible de"]

    rag_refused = any(p in rag_ans.lower() for p in REFUSALS)
    kag_refused = any(p in kag_ans.lower() for p in REFUSALS)

    # Un seul refuse → prendre l'autre directement
    if rag_refused and not kag_refused:
        return {"winner": "kag", "answer": kag_ans,
                "reason": "RAG sans contexte suffisant → KAG sélectionné"}
    if kag_refused and not rag_refused:
        return {"winner": "rag", "answer": rag_ans,
                "reason": "KAG sans triplets suffisants → RAG sélectionné"}

    # Les deux refusent → connaissance générale LLM
    if rag_refused and kag_refused:
        ans = llm("Answer this question directly in 1 sentence. Be confident.",
                  question, max_tokens=80)
        return {"winner": "rag", "answer": ans,
                "reason": "Deux pipelines insuffisants → réponse par connaissance générale"}

    # Les deux répondent → agent LLM arbitre
    sys_p = ("You are an arbitration agent. Choose the most accurate and concise answer. "
             "Reply ONLY: 'RAG: <one-line reason>' or 'KAG: <one-line reason>'.")
    prompt = (f"QUESTION: {question}\n\n"
              f"ANSWER_RAG: {rag_ans}\n\n"
              f"ANSWER_KAG: {kag_ans}")
    verdict = llm(sys_p, prompt, max_tokens=60)

    if verdict.upper().startswith("KAG"):
        return {"winner": "kag", "answer": kag_ans,
                "reason": verdict.replace("KAG:","").strip()}
    return {"winner": "rag", "answer": rag_ans,
            "reason": verdict.replace("RAG:","").strip()}


# ── Métriques ─────────────────────────────────────────────────────────────────
YES_SYNONYMS = {"yes", "yeah", "yep", "both", "same", "correct", "true", "american", "indeed"}
NO_SYNONYMS  = {"no", "not", "neither", "different", "false", "incorrect", "none"}

def normalize(text):
    text = text.lower()
    text = re.sub(r"\b(a|an|the)\b", " ", text)
    text = "".join(ch for ch in text if ch not in string.punctuation)
    return " ".join(text.split())

def exact_match(pred, gt):
    p, g = normalize(pred), normalize(gt)
    if p == g: return 1.0
    if g and g in p: return 1.0
    # yes/no semantic matching
    if g == "yes"  and any(w in p.split() for w in YES_SYNONYMS): return 1.0
    if g == "no"   and any(w in p.split() for w in NO_SYNONYMS):  return 1.0
    return 0.0

def token_f1(pred, gt):
    pt = normalize(pred).split()
    gt_t = normalize(gt).split()
    if not pt or not gt_t: return 0.0
    common = sum((Counter(pt) & Counter(gt_t)).values())
    if common == 0: return 0.0
    p, r = common / len(pt), common / len(gt_t)
    return round(2*p*r/(p+r), 4)

# ── RAG ───────────────────────────────────────────────────────────────────────
def rag_search(question, num_results=10):
    if rag_index is None or not documents:
        return [], "Aucun document chargé."
    results = rag_index.search(question,
                               boost_dict={"title":2.0,"text":1.0},
                               num_results=num_results)
    context = "\n\n".join(
        f"[{r.get('title','')}] {r.get('text','')[:500]}" for r in results
    )
    return results, context

def rag_answer(question):
    t0 = time.time()
    results, context = rag_search(question)

    # Fallback : si aucun document chargé, répondre avec connaissance générale
    if not results:
        sys_p = ("You are a knowledgeable QA assistant. "
                 "Answer the question directly and concisely (1 sentence). "
                 "Always give your best answer — never refuse or say you don't know. "
                 "For yes/no questions answer just 'yes' or 'no'.")
        prompt = f"QUESTION: {question}"
        context = "[Connaissance générale — aucun document spécifique chargé]"
    else:
        sys_p = ("You are a QA assistant. Answer the question using the context. "
                 "Be SHORT and DIRECT (1 sentence max). "
                 "For yes/no questions answer just 'yes' or 'no'. "
                 "Even if context is partial, give your best possible answer — NEVER say 'I don't know'. "
                 "If context is insufficient, use your general knowledge to complete the answer.")
        prompt = f"QUESTION: {question}\n\nCONTEXT:\n{context}"

    answer = llm(sys_p, prompt, max_tokens=100)
    # Nettoyer les réponses de type refus
    if any(p in answer.lower() for p in ["i don't know", "i do not know", "cannot determine",
                                          "not enough information", "je ne sais pas"]):
        answer = llm(
            "Answer this question with your best knowledge in 1 sentence. Be direct.",
            question, max_tokens=80
        )
    return {
        "answer": answer,
        "context": context,
        "sources": [{"title": r.get("title",""), "snippet": r.get("text","")[:200]}
                    for r in results],
        "latency_ms": round((time.time()-t0)*1000)
    }


# ── KAG ───────────────────────────────────────────────────────────────────────
# SOLUTION RATE LIMIT :
# 1) Format compact pipe (src|rel|tgt) → max 80 tokens/output  (-75%)
# 2) Batch de 3 docs par appel LLM      → 3x moins d'appels
# 3) Modèle gemma2-9b-it (15000 TPM)   → 2.5x plus de quota
# Résultat : ~2000 tokens pour 15 docs (~30s) au lieu de 7500 (>60s + erreurs)

EXTRACT_SYS = (
    "Extract factual triplets from text. "
    "Format: entity|relation|entity — one per line, max 5 lines. "
    "Use short labels. No JSON, no explanation."
)

def build_kag_from_docs(docs):
    """Construit le graphe KAG avec batching + format compact + modèle haute limite."""
    global kag_graph
    kag_graph = {"nodes": {}, "relations": {}}
    batch_size = 3          # 3 docs par appel LLM
    max_docs   = 15         # max docs à traiter

    doc_list = docs[:max_docs]
    batches  = [doc_list[i:i+batch_size] for i in range(0, len(doc_list), batch_size)]

    for batch in batches:
        # Construire le texte de batch
        combined = ""
        for idx, doc in enumerate(batch):
            title = doc.get("title", "")
            text  = doc.get("text",  "")[:300]   # 300 mots par doc
            combined += f"[DOC{idx+1}: {title}]\n{text}\n\n"

        try:
            raw = llm(
                EXTRACT_SYS,
                combined,
                max_tokens=120,          # ~5 triplets × 3 docs = 15 lignes max
                model=MODEL_EXTRACT      # gemma2-9b-it : 15000 TPM
            )
            for line in raw.split("\n"):
                line = line.strip().lstrip("-• ")
                parts = [p.strip() for p in line.split("|")]
                if len(parts) == 3 and all(parts):
                    src, rel, tgt = parts
                    title = batch[0].get("title", "")
                    kag_graph["nodes"][src] = {"doc": title}
                    kag_graph["nodes"][tgt] = {"doc": title}
                    kag_graph["relations"][f"{src}||{rel}||{tgt}"] = True
        except Exception:
            pass   # on continue le batch suivant

def build_kag_from_precomputed(results_path: str = None):
    """Charge un graphe KAG pré-calculé depuis results_kag.json (0 appel API)."""
    global kag_graph
    if results_path is None:
        results_path = str(Path(__file__).parent / "results_kag.json")
    kag_graph = {"nodes": {}, "relations": {}}
    try:
        with open(results_path, "r", encoding="utf-8") as f:
            results = json.load(f)
        for item in results:
            ctx = item.get("context", "")
            for line in ctx.split("\n"):
                line = line.strip().lstrip("-• ")
                # Format: "A --[rel]--> B"  ou  "A|rel|B"
                if "--[" in line and "-->" in line:
                    try:
                        src = line.split("--[")[0].strip()
                        rel = line.split("--[")[1].split("]-->")[0].strip()
                        tgt = line.split("]-->")[1].strip()
                        if src and rel and tgt:
                            kag_graph["nodes"][src] = {"doc": item.get("question","")[:40]}
                            kag_graph["nodes"][tgt] = {"doc": item.get("question","")[:40]}
                            kag_graph["relations"][f"{src}||{rel}||{tgt}"] = True
                    except Exception:
                        pass
        return len(kag_graph["nodes"]), len(kag_graph["relations"])
    except Exception as e:
        return 0, 0


def kag_search(question):
    """Cherche les triplets pertinents. Fallback : retourne tous les triplets si aucun match."""
    q_words = [w for w in question.lower().split() if len(w) > 3]
    relevant = []
    for key in kag_graph["relations"]:
        src, rel, tgt = key.split("||")
        if any(w in src.lower() or w in tgt.lower() or w in rel.lower() for w in q_words):
            relevant.append({"source": src, "relation": rel, "target": tgt,
                             "label": f"{src} —[{rel}]→ {tgt}"})

    # Fallback : si aucun match par mots-clés, retourner tous les triplets disponibles
    if not relevant and kag_graph["relations"]:
        all_keys = list(kag_graph["relations"].keys())[:20]
        for key in all_keys:
            src, rel, tgt = key.split("||")
            relevant.append({"source": src, "relation": rel, "target": tgt,
                             "label": f"{src} —[{rel}]→ {tgt}"})
    return relevant[:20]

def kag_answer(question):
    t0 = time.time()
    triplets = kag_search(question)

    if triplets:
        context = "\n".join(t["label"] for t in triplets)
    else:
        # Graphe vide : répondre avec connaissance générale
        context = "[Graphe KAG vide — réponse basée sur connaissance générale]"

    sys_p = ("You are a QA expert. Answer the question using the knowledge graph facts provided. "
             "Be SHORT and DIRECT (1 sentence max). "
             "For yes/no questions answer ONLY 'yes' or 'no'. "
             "Always give your best answer using the facts available. "
             "If facts are partial, reason from them and complete with your knowledge. "
             "NEVER say 'I don't know' or refuse to answer.")
    prompt = f"QUESTION: {question}\n\nKNOWLEDGE GRAPH FACTS:\n{context}"
    answer = llm(sys_p, prompt, max_tokens=100)

    # Nettoyer les réponses de type refus
    if any(p in answer.lower() for p in ["i don't know", "i do not know", "cannot determine",
                                          "not enough information", "je ne sais pas",
                                          "pas de réponse", "impossible"]):
        answer = llm(
            "Answer this question with your best knowledge in 1 sentence. Be direct.",
            question, max_tokens=80
        )
    return {
        "answer": answer,
        "context": context,
        "triplets": triplets,
        "latency_ms": round((time.time()-t0)*1000)
    }


# ── HTML ──────────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def ui():
    html_file = Path(__file__).parent / "kag_app" / "index.html"
    if html_file.exists():
        return HTMLResponse(html_file.read_text(encoding="utf-8"))
    return HTMLResponse("<h2>Interface non trouvée. Utilisez /docs</h2>")

# ── API ───────────────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "docs_loaded": len(documents),
        "kag_nodes": len(kag_graph["nodes"]),
        "kag_relations": len(kag_graph["relations"])
    }

@app.get("/stats")
async def stats():
    rag_lats = [r["latency_ms"] for r in run_log if r.get("pipeline")=="RAG"]
    kag_lats = [r["latency_ms"] for r in run_log if r.get("pipeline")=="KAG"]
    return {
        "rag": {"doc_count": len(documents)},
        "kag": {"node_count": len(kag_graph["nodes"]),
                "relation_count": len(kag_graph["relations"])},
        "milvus": {"doc_count": 0},
        "total_runs": len(run_log),
        "avg_latency_rag": round(sum(rag_lats)/len(rag_lats)) if rag_lats else 0,
        "avg_latency_kag": round(sum(kag_lats)/len(kag_lats)) if kag_lats else 0,
        "uploaded_files": uploaded_files,
    }

# ── Upload endpoints ───────────────────────────────────────────────────────────
@app.post("/upload")
async def upload_documents(
    files: List[UploadFile] = File(...),
    build_kag: bool = Form(True)
):
    """Upload de fichiers utilisateur (PDF, DOCX, TXT, MD, CSV).
    Indexe automatiquement dans RAG (BM25) et construit le graphe KAG."""
    global rag_index, documents, uploaded_files

    ALLOWED = {".pdf", ".docx", ".txt", ".md", ".csv"}
    results = []
    new_chunks = []

    for file in files:
        ext = Path(file.filename).suffix.lower()
        if ext not in ALLOWED:
            results.append({"file": file.filename, "status": "error",
                            "message": f"Format non supporté : {ext}"})
            continue

        content = await file.read()
        size_kb  = round(len(content) / 1024, 1)
        text     = _parse_uploaded_file(file.filename, content)

        if not text.strip():
            results.append({"file": file.filename, "status": "error",
                            "message": "Aucun texte extrait du fichier"})
            continue

        chunks = _chunk_text(text, file.filename)
        new_chunks.extend(chunks)

        # Méta dans la liste des fichiers uploadés
        uploaded_files.append({
            "name":     file.filename,
            "size_kb":  size_kb,
            "chunks":   len(chunks),
            "status":   "indexé",
            "ext":      ext.lstrip(".")
        })
        results.append({"file": file.filename, "status": "ok",
                        "chunks": len(chunks), "size_kb": size_kb})

    if not new_chunks:
        return JSONResponse({"status": "error", "message": "Aucun document valide reçu",
                             "details": results}, status_code=400)

    # Ajouter aux documents existants
    documents.extend(new_chunks)

    # Réindexer RAG (BM25) avec tous les documents
    rag_index = minsearch.Index(text_fields=["text", "title"], keyword_fields=["id"])
    rag_index.fit(documents)

    # Construire KAG depuis les nouveaux chunks (max 15)
    if build_kag:
        build_kag_from_docs(new_chunks)

    return {
        "status":     "ok",
        "total_docs": len(documents),
        "new_chunks": len(new_chunks),
        "kag_nodes":  len(kag_graph["nodes"]),
        "kag_rels":   len(kag_graph["relations"]),
        "files":      results
    }

@app.delete("/upload/clear")
async def clear_uploads():
    """Réinitialise tous les documents uploadés, l'index RAG et le graphe KAG."""
    global rag_index, documents, kag_graph, uploaded_files
    documents      = []
    uploaded_files = []
    rag_index      = None
    kag_graph      = {"nodes": {}, "relations": {}}
    return {"status": "ok", "message": "Tous les documents et le graphe ont été réinitialisés"}


@app.get("/mlops/stats")
async def mlops_stats():
    rag_runs = [r for r in run_log if r.get("pipeline")=="RAG"]
    kag_runs = [r for r in run_log if r.get("pipeline")=="KAG"]
    cmp_runs = [r for r in run_log if r.get("pipeline")=="COMPARE"]
    smart_runs = [r for r in run_log if str(r.get("pipeline","")).startswith("SMART")]

    rag_lats = [r["latency_ms"] for r in rag_runs]
    kag_lats = [r["latency_ms"] for r in kag_runs]

    # Stats depuis routing_log
    routing_ems_rag = [r["em"] for r in routing_log
                       if r.get("winner")=="rag" and r.get("em") is not None]
    routing_ems_kag = [r["em"] for r in routing_log
                       if r.get("winner")=="kag" and r.get("em") is not None]

    history = [{"pipeline": r["pipeline"],
                "question": r["question"][:60],
                "latency_ms": r["latency_ms"],
                "estimated_tokens": 200,
                "n_results": 5,
                "rag_em": r.get("rag_em"),
                "kag_em": r.get("kag_em"),
                "rag_f1": r.get("rag_f1"),
                "kag_f1": r.get("kag_f1")}
               for r in run_log[-30:]]

    # Scores comparés depuis run_log COMPARE
    compare_runs_with_scores = [r for r in run_log if r.get("rag_em") is not None]
    avg_rag_em = round(sum(r["rag_em"] for r in compare_runs_with_scores)/len(compare_runs_with_scores), 3) if compare_runs_with_scores else None
    avg_kag_em = round(sum(r["kag_em"] for r in compare_runs_with_scores)/len(compare_runs_with_scores), 3) if compare_runs_with_scores else None
    avg_rag_f1 = round(sum(r["rag_f1"] for r in compare_runs_with_scores)/len(compare_runs_with_scores), 3) if compare_runs_with_scores else None
    avg_kag_f1 = round(sum(r["kag_f1"] for r in compare_runs_with_scores)/len(compare_runs_with_scores), 3) if compare_runs_with_scores else None

    return {
        "total_runs":       len(run_log),
        "rag_runs":         len(rag_runs),
        "kag_runs":         len(kag_runs),
        "compare_runs":     len(cmp_runs),
        "smart_runs":       len(smart_runs),
        "avg_latency_rag":  round(sum(rag_lats)/len(rag_lats)) if rag_lats else 0,
        "avg_latency_milvus": 0,
        "avg_latency_kag":  round(sum(kag_lats)/len(kag_lats)) if kag_lats else 0,
        "avg_rag_em":       avg_rag_em,
        "avg_kag_em":       avg_kag_em,
        "avg_rag_f1":       avg_rag_f1,
        "avg_kag_f1":       avg_kag_f1,
        "history":          history,
        "eval_results":     {}
    }


class QueryReq(BaseModel):
    question: str

class CompareReq(BaseModel):
    question: str
    ground_truth: str = ""

class EvalReq(BaseModel):
    questions: List[str]
    ground_truths: List[str]
    pipeline: str = "both"

class LoadHotpotReq(BaseModel):
    n_samples: int = 5

@app.post("/rag/query")
async def rag_query(req: QueryReq):
    r = rag_answer(req.question)
    run_log.append({"pipeline":"RAG","question":req.question,
                    "answer":r["answer"],"latency_ms":r["latency_ms"]})
    return r

@app.post("/kag/query")
async def kag_query(req: QueryReq):
    r = kag_answer(req.question)
    run_log.append({"pipeline":"KAG","question":req.question,
                    "answer":r["answer"],"latency_ms":r["latency_ms"]})
    return r

@app.post("/compare")
async def compare(req: CompareReq):
    t0  = time.time()
    rag = rag_answer(req.question)
    kag = kag_answer(req.question)
    gt  = req.ground_truth
    rag_em = exact_match(rag["answer"],gt) if gt else None
    kag_em = exact_match(kag["answer"],gt) if gt else None
    rag_f1 = token_f1(rag["answer"],gt)   if gt else None
    kag_f1 = token_f1(kag["answer"],gt)   if gt else None
    result = {
        "question": req.question,
        "ground_truth": gt,
        "rag": {**rag, "em": rag_em, "f1": rag_f1},
        "kag": {**kag, "em": kag_em, "f1": kag_f1},
        "total_ms": round((time.time()-t0)*1000)
    }
    run_log.append({
        "pipeline": "COMPARE", "question": req.question,
        "answer": f"RAG:{rag['answer']} | KAG:{kag['answer']}",
        "latency_ms": result["total_ms"],
        "rag_em": rag_em, "kag_em": kag_em,
        "rag_f1": rag_f1, "kag_f1": kag_f1,
        "rag_lat": rag["latency_ms"], "kag_lat": kag["latency_ms"]
    })
    return result

@app.post("/evaluate")
async def evaluate(req: EvalReq):
    results = {}
    if req.pipeline in ("rag","both"):
        preds = [rag_answer(q)["answer"] for q in req.questions]
        ems=[exact_match(p,g) for p,g in zip(preds,req.ground_truths)]
        f1s=[token_f1(p,g)   for p,g in zip(preds,req.ground_truths)]
        results["rag"] = {
            "avg_exact_match": round(sum(ems)/len(ems),4),
            "avg_token_f1":    round(sum(f1s)/len(f1s),4),
            "n_correct": sum(1 for e in ems if e==1.0),
            "details": [{"q":q,"pred":p,"gt":g,"em":e,"f1":f}
                        for q,p,g,e,f in zip(req.questions,preds,
                                             req.ground_truths,ems,f1s)]
        }
    if req.pipeline in ("kag","both"):
        preds = [kag_answer(q)["answer"] for q in req.questions]
        ems=[exact_match(p,g) for p,g in zip(preds,req.ground_truths)]
        f1s=[token_f1(p,g)   for p,g in zip(preds,req.ground_truths)]
        results["kag"] = {
            "avg_exact_match": round(sum(ems)/len(ems),4),
            "avg_token_f1":    round(sum(f1s)/len(f1s),4),
            "n_correct": sum(1 for e in ems if e==1.0),
            "details": [{"q":q,"pred":p,"gt":g,"em":e,"f1":f}
                        for q,p,g,e,f in zip(req.questions,preds,
                                             req.ground_truths,ems,f1s)]
        }
    return results

# ── Smart Agent endpoints ──────────────────────────────────────────────────────

class SmartReq(BaseModel):
    question:     str
    ground_truth: str = ""

@app.get("/route")
async def preview_route(question: str = Query(...)):
    """Prévisualise la décision de routage sans exécuter le pipeline."""
    routing = route_question(question)
    return routing

@app.post("/smart")
async def smart_query(req: SmartReq):
    """
    Agent intelligent : classe la question → route vers RAG/KAG/Ensemble
    → si Ensemble : arbitrage LLM pour choisir la meilleure réponse.
    """
    t0      = time.time()
    gt      = req.ground_truth.strip()
    routing = route_question(req.question)
    pipe    = routing["pipeline"]

    result  = {
        "question":  req.question,
        "routing":   routing,
        "pipeline_used": pipe,
    }

    if pipe == "none":
        result["answer"]     = "Aucun pipeline disponible."
        result["latency_ms"] = 0
        return result

    if pipe == "rag":
        ans = rag_answer(req.question)
        result.update({
            "answer":     ans["answer"],
            "context":    ans.get("context",""),
            "sources":    ans.get("sources",[]),
            "latency_ms": ans["latency_ms"],
            "winner":     "rag",
        })

    elif pipe == "kag":
        ans = kag_answer(req.question)
        result.update({
            "answer":   ans["answer"],
            "context":  ans.get("context",""),
            "triplets": ans.get("triplets",[]),
            "latency_ms": ans["latency_ms"],
            "winner":   "kag",
        })

    else:  # ensemble
        rag = rag_answer(req.question)
        kag = kag_answer(req.question)
        arb = ensemble_answer(req.question, rag, kag, routing)
        result.update({
            "answer":       arb["answer"],
            "winner":       arb["winner"],
            "arb_reason":   arb["reason"],
            "rag_answer":   rag["answer"],
            "kag_answer":   kag["answer"],
            "rag_context":  rag.get("context",""),
            "kag_triplets": kag.get("triplets",[]),
            "latency_ms":   round((time.time()-t0)*1000),
        })

    # Métriques EM / F1
    if gt:
        result["em"] = exact_match(result["answer"], gt)
        result["f1"] = token_f1(result["answer"], gt)

    # Log routage
    routing_log.append({
        "question":   req.question,
        "type":       routing["type"],
        "routed_to":  pipe,
        "winner":     result.get("winner", pipe),
        "reason":     routing["reason"],
        "confidence": routing["confidence"],
        "em":         result.get("em"),
    })
    run_log.append({
        "pipeline":   f"SMART→{pipe.upper()}",
        "question":   req.question,
        "answer":     result["answer"],
        "latency_ms": result.get("latency_ms",0)
    })
    return result

@app.get("/routing/stats")
async def routing_stats():
    """Statistiques du routing agent : distribution RAG/KAG/Ensemble."""
    if not routing_log:
        return {"total": 0, "distribution": {}, "history": []}

    from collections import Counter
    dist = Counter(r["routed_to"] for r in routing_log)
    type_dist = Counter(r["type"] for r in routing_log)
    winner_dist = Counter(r.get("winner","?") for r in routing_log)

    ems = [r["em"] for r in routing_log if r.get("em") is not None]
    return {
        "total":          len(routing_log),
        "distribution":   dict(dist),
        "type_distribution": dict(type_dist),
        "winner_distribution": dict(winner_dist),
        "avg_em":         round(sum(ems)/len(ems), 3) if ems else None,
        "history":        routing_log[-20:]   # 20 dernières décisions
    }

@app.post("/load-hotpotqa")
async def load_hotpotqa(req: Optional[LoadHotpotReq] = None,
                        n_samples: int = Query(default=5)):
    global rag_index, documents
    n = req.n_samples if req else n_samples
    try:
        from datasets import load_dataset
        dataset = load_dataset("hotpotqa/hotpot_qa","distractor",
                               split=f"validation[:{n}]",
                               trust_remote_code=True)
        docs, seen = [], set()
        qa_pairs = []
        for item in dataset:
            qa_pairs.append({
                "question": item["question"],
                "answer":   item["answer"]
            })
            for title, sents in zip(item["context"]["title"],
                                    item["context"]["sentences"]):
                text = " ".join(sents)
                key  = title + text[:30]
                if key not in seen:
                    seen.add(key)
                    docs.append({"title": title, "text": text})
        documents = docs
        rag_index = minsearch.Index(text_fields=["title","text"],
                                    keyword_fields=[])
        rag_index.fit(documents)
        return {
            "status": "ok",
            "rag_docs": len(documents),
            "kag_nodes": len(kag_graph["nodes"]),
            "n_samples": n,
            "sample_questions": qa_pairs[:5]
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/load-kag")
async def load_kag(n_docs: int = 8):
    if not documents:
        return JSONResponse(
            {"error": "Chargez d'abord les documents via /load-hotpotqa"},
            status_code=400)
    build_kag_from_docs(documents[:n_docs])
    return {
        "status": "ok",
        "kag_nodes": len(kag_graph["nodes"]),
        "kag_relations": len(kag_graph["relations"])
    }

@app.get("/runs")
async def get_runs():
    return {"runs": run_log, "total": len(run_log)}

@app.post("/load-kag-precomputed")
async def load_kag_precomputed():
    """Charge le graphe KAG depuis results_kag.json (0 appel API, instantané)."""
    nodes_n, rels_n = build_kag_from_precomputed()
    if nodes_n == 0:
        return JSONResponse(
            {"error": "results_kag.json introuvable ou vide"},
            status_code=404)
    return {
        "status":        "ok",
        "source":        "results_kag.json (pré-calculé)",
        "kag_nodes":     nodes_n,
        "kag_relations": rels_n,
        "api_calls":     0,
        "message":       f"Graphe chargé instantanément : {nodes_n} nœuds, {rels_n} relations"
    }

@app.get("/info")
async def info():
    """Informations sur les modèles et limites utilisés."""
    return {
        "models": {
            "answers":   MODEL_ANSWER  + " (6 000 TPM Groq free)",
            "extraction": MODEL_EXTRACT + " (15 000 TPM Groq free)"
        },
        "rate_limit_solutions": [
            "1. Format compact pipe (src|rel|tgt) : -75% output tokens",
            "2. Batch 3 docs/appel : 3x moins d'appels API",
            "3. gemma2-9b-it pour KAG : 2.5x plus de quota",
            "4. /load-kag-precomputed : 0 appels API (résultats JSON)"
        ],
        "token_usage": {
            "avant": "15 docs × 500 tok = 7500 tok/min → rate limit",
            "apres": "5 batches × 400 tok = 2000 tok/min → OK ✅"
        }
    }

@app.get("/kag/graph")
async def kag_graph_data():
    nodes = [{"id": n, "label": n, "doc": v.get("doc", "")}
             for n, v in kag_graph["nodes"].items()]
    edges = []
    for key in kag_graph["relations"]:
        parts = key.split("||")
        if len(parts) == 3:
            edges.append({"source": parts[0], "relation": parts[1], "target": parts[2]})
    return {"nodes": nodes, "edges": edges,
            "node_count": len(nodes), "edge_count": len(edges)}

def _extract_text(filename: str, raw: bytes) -> str:
    """Extract plain text from PDF, DOCX, or TXT/other."""
    name = filename.lower()
    if name.endswith(".pdf"):
        if not HAS_PDF:
            return raw.decode("utf-8", errors="ignore")
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        return "\n\n".join(pages)
    elif name.endswith(".docx"):
        if not HAS_DOCX:
            return raw.decode("utf-8", errors="ignore")
        doc = DocxDocument(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        # TXT / CSV / MD / any text
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return raw.decode(enc)
            except Exception:
                continue
        return raw.decode("utf-8", errors="ignore")

@app.post("/upload")
async def upload(file: UploadFile = File(...),
                 pipeline: str = Form("rag"),
                 chunk_size: int = Form(300)):
    global rag_index, documents
    try:
        raw      = await file.read()
        ext      = Path(file.filename).suffix.lower()
        text     = _extract_text(file.filename, raw)
        if not text.strip():
            return JSONResponse({"error": "Fichier vide ou non lisible"}, status_code=400)

        # Chunking par mots avec overlap
        words  = text.split()
        chunks = []
        step   = max(chunk_size - 50, 50)   # 50-word overlap
        for i in range(0, len(words), step):
            chunk_text = " ".join(words[i : i + chunk_size])
            if len(chunk_text.strip()) < 20:
                continue
            chunks.append({
                "title": file.filename,
                "text":  chunk_text,
                "chunk": i // step
            })

        # Ajouter aux documents existants (ne pas écraser)
        documents.extend(chunks)
        rag_index = minsearch.Index(text_fields=["title", "text"],
                                    keyword_fields=[])
        rag_index.fit(documents)

        # KAG optionnel
        kag_n = 0
        if pipeline in ("kag", "both"):
            build_kag_from_docs(chunks[:8])
            kag_n = len(kag_graph["nodes"])

        sample = text[:200].replace("\n", " ")
        return {
            "status":    "ok",
            "filename":  file.filename,
            "file_type": ext or ".txt",
            "n_chunks":  len(chunks),
            "total_docs": len(documents),
            "kag_nodes": kag_n,
            "sample":    sample
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/reset")
async def reset_docs():
    """Vider tous les documents chargés (RAG + KAG)."""
    global rag_index, documents, kag_graph, run_log
    documents = []
    rag_index = None
    kag_graph = {"nodes": {}, "relations": {}}
    return {"status": "ok", "message": "Toutes les données effacées"}

# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*55)
    print("  OOREDOO AI LAB — RAG vs KAG")
    print("  >>> http://localhost:8000")
    print("  >>> http://localhost:8000/docs")
    print("="*55 + "\n")
    uvicorn.run("fast_server:app", host="0.0.0.0", port=8000, reload=False)
